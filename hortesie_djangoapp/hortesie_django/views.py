import json
import os
import time
import uuid

from django.core.cache import cache
from django.http import HttpResponse
from django.templatetags.static import static
from docx import Document
from rest_framework import status
from rest_framework.decorators import action
from rest_framework.parsers import MultiPartParser
from rest_framework.response import Response
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework.views import APIView
from rest_framework.viewsets import ModelViewSet

from hortesie_django.models import CCTPTemplate, Photo, Project, Article
from hortesie_django.serializers import (
	CCTPTemplateSerializer,
	PhotoSerializer,
	ProjectSerializer,
	PhotoModelSerializer,
	ArticleSerializer,
)


def generate_uid():
	return str(uuid.uuid1())


def extract_headings_with_content(doc):
	headings = []
	current_heading = None
	heading_stack = []
	compteur_part_principale = 0
	level_labels = [0, 0, 0, 0, 0, 0]
	for paragraph in doc.paragraphs:
		if "Heading" in paragraph.style.name:
			# Determine the depth of the current heading based on indentation
			# Pop headings from the stack until we find the parent of the current heading
			# Set the parent heading
			heading_uid = generate_uid()
			# If a new heading is found, save the previous heading and its content
			if current_heading is not None:
				headings.append(current_heading)
			last_heading = ""
			heading_level = int(paragraph.style.name[7:])
			if len(heading_stack) > 0:
				last_heading = heading_stack[-1]
			if last_heading != "" and heading_level < last_heading["depth"]:
				level_labels[heading_level - 1] += 1
				for i in range(heading_level, last_heading["depth"]):
					level_labels[i] = 0
			else:
				level_labels[heading_level - 1] += 1
			if heading_level == 1:
				compteur_part_principale += 1
			current_heading = {
				"uid": heading_uid,
				"title": paragraph.text,
				"content": [],
				"depth": heading_level,
				"label_part": ".".join(
					[
						str(level_labels[i])
						for i in range(len(level_labels))
						if i < heading_level
					]
				),
				"isChecked": True,
				"isChildDisplayed": False,
				"isDisplayed": True,
			}
			# Push the current heading to the stack
			heading_stack.append(current_heading)

		elif current_heading is not None:
			# Append the content of the current heading
			current_heading["content"].append(paragraph.text)

	# Append the last heading
	if current_heading is not None:
		current_heading["parent"] = (
			heading_stack[-2] if len(heading_stack) >= 2 else None
		)
		headings.append(current_heading)
	return headings


def get_modified_date(file_path):
	try:
		# Get the modification timestamp of the file in seconds since the epoch
		modification_timestamp = os.path.getmtime(file_path)

		# Convert the timestamp to a human-readable date format
		modified_date = time.strftime(
			"%Y-%m-%d %H:%M:%S", time.localtime(modification_timestamp)
		)

		return modified_date

	except FileNotFoundError:
		return None


def get_structure(doc):
	headings = []
	heading_stack = []
	current_heading = None
	for idx, paragraph in enumerate(doc.paragraphs):

		if "Heading" in paragraph.style.name:
			# Set the parent heading
			parent_heading = heading_stack[-1] if heading_stack else None

			# If a new heading is found, save the previous heading  and its content
			if current_heading is not None:
				current_heading["parent"] = parent_heading
				headings.append(current_heading)

			# Generate a UID for the current paragraph
			current_heading = {
				"title": paragraph.text,
				"content": [],
				"depth": int(paragraph.style.name[7:]),
				"paragraph": [paragraph],
			}

			# Push the current heading to the stack
			heading_stack.append(current_heading)

		elif current_heading is not None:
			# Append the content of the current heading
			current_heading["content"].append(paragraph.text)
			current_heading["paragraph"].append(paragraph)

	# Append the last heading
	if current_heading is not None:
		current_heading["parent"] = (
			heading_stack[-2] if len(heading_stack) >= 2 else None
		)
		headings.append(current_heading)
	return headings


def delete_all_title(doc, match_dict):
	headings = get_structure(doc)
	index = 0
	while index < len(headings):
		heading = headings[index]
		if any(
				heading["title"] == item.get("title")
				and heading["content"] == item.get("content")
				for item in match_dict
		):
			for parag in heading["paragraph"]:
				p = parag._element
				p.getparent().remove(p)
				p._element = None
			del headings[index]
			index -= 1
		index += 1
	return headings


def remove_heading(doc, match_dict):
	delete_all_title(doc, match_dict)

	doc.save(r"hortesie_django/files/TEMPLATE_WORD_CCTP_temp.docx")


class GenerateCSVAPIView(APIView):

	def get_last_cctp_file_url(self):
		qs = CCTPTemplate.objects.order_by("-date")
		if qs.exists():
			return qs[0].file.url.replace("/media", "files")

	def get(self, request):
		# Generate the CSV content in memory using io.BytesIO
		cached_data = cache.get("CCTP_structure")
		url = self.get_last_cctp_file_url()
		print("url iss",url)
		date_modif = get_modified_date(url)
		doc = Document(url)
		if not cache.get("CCTP_date") or date_modif > cache.get("CCTP_date"):
			# Data is not in the cache, compute it
			doc = Document(url)
			data = extract_headings_with_content(doc)
			# Store the data in the cache for future use
			cache.set("CCTP_structure", data)
			cache.set("CCTP_date", date_modif)

		cached_data = cache.get("CCTP_structure")
		# Set the appropriate response headers
		response = Response(cached_data)
		return response

	def post(self, request):
		url = self.get_last_cctp_file_url()
		values = json.loads(request.body.decode("utf-8")).get("values")
		filename = json.loads(request.body.decode("utf-8")).get("filename")
		doc = Document(url)
		remove_heading(doc, values)
		file_path = "hortesie_django/files/TEMPLATE_WORD_CCTP_temp.docx"
		with open(file_path, "rb") as file:
			doc_data = file.read()
		# sending response
		response = HttpResponse(
			doc_data,
			content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
		)
		response["Content-Disposition"] = f'attachment; filename="{filename}.docx"'

		return response


class ProjectViewset(ModelViewSet):
    queryset = Project.objects.all()
    serializer_class = ProjectSerializer

    @action(detail=True, methods=["post"])
    def add_image(self, request, pk):
        project = self.get_object()
        serializer = PhotoSerializer(data=request.data, instance=project)

        if serializer.is_valid():
                serializer.create(serializer.validated_data)

                return Response(serializer.data, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def retrieve(self, request, *args, **kwargs):
            serializer = self.get_serializer(self.get_object())
            instance = self.get_object()
            if instance.vignette:
                    return Response(
                            {
                                    **serializer.data,
                                    "vignette": "/api" + self.get_object().vignette.file.url,
                            }
                    )
            return Response(serializer.data)

    @action(detail=True, methods=["get"])
    def images(self, request, pk):
            project = self.get_object()
            data = Photo.objects.filter(id_projet=project.id)
            print(data)
            serializer = PhotoModelSerializer(data, many=True)
            print(serializer.data)
            return Response(serializer.data)

    @action(detail=False, methods=["post"])
    def switch(self, request):
            from_index = request.data.get("from")
            to_index = request.data.get("to")
            a = Project.objects.get(position=from_index)

            b = Project.objects.get(position=to_index)
            b.position = from_index
            a.position = to_index
            a.save()
            b.save()
            return Response(status=status.HTTP_200_OK)

    @action(detail=False, methods=["post"])
    def update_position(self, request):
            """
            Update a project's position directly.
            Expects: {"project_id": "id", "new_position": number}
            """
            project_id = request.data.get("project_id")
            new_position = request.data.get("new_position")
            
            if not project_id or new_position is None:
                    return Response(
                            {"error": "project_id and new_position are required"}, 
                            status=status.HTTP_400_BAD_REQUEST
                    )
            
            try:
                    project = Project.objects.get(id=project_id)
                    old_position = project.position
                    
                    # Update the project's position
                    project.position = new_position
                    project.save()
                    
                    return Response({
                            "success": True,
                            "project_id": project_id,
                            "old_position": old_position,
                            "new_position": new_position
                    }, status=status.HTTP_200_OK)
                    
            except Project.DoesNotExist:
                    return Response(
                            {"error": "Project not found"}, 
                            status=status.HTTP_404_NOT_FOUND
                    )
            except Exception as e:
                    return Response(
                            {"error": str(e)}, 
                            status=status.HTTP_500_INTERNAL_SERVER_ERROR
                    )

    def list(self, request, *args, **kwargs):
            data = super().list(request, *args, **kwargs)
            data = sorted(data.data, key=lambda d: d["position"])

            return Response(data)


class CCTPTemplateCreateView(APIView):
	parser_classes = [MultiPartParser]  # This allows file uploads in the request body

	def post(self, request, *args, **kwargs):
		# Initialize the serializer with the request data
		serializer = CCTPTemplateSerializer(data=request.data)

		if serializer.is_valid():
			# Save the instance and return the serialized response
			cctp_template = serializer.create(serializer.validated_data)
			return Response(
				{
					"id": cctp_template.id,
					"file": cctp_template.file.url,
				},
				status=status.HTTP_201_CREATED,
			)
		return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class PhotoViewSet(ModelViewSet):
	queryset = Photo.objects.all()


class ArticleViewset(ModelViewSet):
    queryset = Article.objects.all()
    serializer_class = ArticleSerializer

    @action(detail=True, methods=["post"])
    def add_image(self, request, pk):
        article = self.get_object()
        serializer = PhotoSerializer(data=request.data, instance=article)
        if serializer.is_valid():
            serializer.create(serializer.validated_data)
            return Response(serializer.data, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    def retrieve(self, request, *args, **kwargs):
        serializer = self.get_serializer(self.get_object())
        instance = self.get_object()
        data = serializer.data
        # ensure proper URLs for media
        if getattr(instance, "vignette", None):
            data["vignette"] = "/api" + instance.vignette.file.url
        if getattr(instance, "pdf", None):
            data["pdf"] = "/api" + instance.pdf.url
        return Response(data)

    @action(detail=True, methods=["get"])
    def images(self, request, pk):
        article = self.get_object()
        data = Photo.objects.filter(id_projet=article.id)
        serializer = PhotoModelSerializer(data, many=True)
        return Response(serializer.data)
